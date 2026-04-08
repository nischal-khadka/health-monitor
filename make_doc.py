from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def heading3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("NOTE: " + text)
    run.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
    run.font.size = Pt(10)
    run.bold = True
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading = p._element
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(" — " + text)
    else:
        p.add_run(text)
    return p

def spacer():
    doc.add_paragraph("")

# ── Title Page ────────────────────────────────────────────────────────────────
title = doc.add_heading("IoT Patient Health Monitor", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub = doc.add_paragraph("Friend Setup Guide — From Zero to Running")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph("ESP32 + MAX30102 + DS18B20  →  FastAPI Backend  →  React Dashboard").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Overview ──────────────────────────────────────────────────────────────────
heading1("Overview")
body("This guide walks you through setting up the entire health monitoring system on your machine. By the end, your ESP32 will be reading heart rate, SpO2, and temperature — sending it to a backend server — and displaying it live on a dashboard in your browser.")

spacer()
body("What you will need:")
bullet("ESP32 development board (ESP32-WROOM-32)")
bullet("MAX30102 sensor (heart rate + SpO2)")
bullet("DS18B20 waterproof temperature sensor")
bullet("4.7kΩ resistor (for DS18B20)")
bullet("Jumper wires + breadboard")
bullet("USB cable to connect ESP32 to your PC")
bullet("PC running Linux, macOS, or Windows with Python 3.10+ and Node.js 18+")

doc.add_page_break()

# ── PART 1 ────────────────────────────────────────────────────────────────────
heading1("PART 1 — Get the Code")

heading2("Step 1: Install Git (if not installed)")
body("Open your terminal and run:")
code("sudo apt install git -y          # Linux (Debian/Ubuntu/Kali)")
code("# macOS: git comes pre-installed")
code("# Windows: download from https://git-scm.com")

spacer()
heading2("Step 2: Clone the Repository")
body("Run this in your terminal:")
code("git clone git@github.com:nischal-khadka/health-monitor.git")
code("cd health-monitor")
note("If you don't have SSH set up, use HTTPS instead:")
code("git clone https://github.com/nischal-khadka/health-monitor.git")

spacer()
body("You should now see this folder structure:")
code("health-monitor/")
code("├── backend/       ← FastAPI server")
code("├── frontend/      ← React dashboard")
code("├── firmware/      ← ESP32 code")
code("└── README.md")

doc.add_page_break()

# ── PART 2 ────────────────────────────────────────────────────────────────────
heading1("PART 2 — Hardware Setup (ESP32 + Sensors)")

heading2("Step 3: Wire the MAX30102 Sensor (Heart Rate + SpO2)")
body("Connect the MAX30102 to the ESP32 using jumper wires:")

table1 = doc.add_table(rows=5, cols=2)
table1.style = "Table Grid"
hdr = table1.rows[0].cells
hdr[0].text = "MAX30102 Pin"
hdr[1].text = "ESP32 Pin"
rows = [("VIN", "3.3V"), ("GND", "GND"), ("SDA", "GPIO 21"), ("SCL", "GPIO 22")]
for i, (a, b) in enumerate(rows, 1):
    table1.rows[i].cells[0].text = a
    table1.rows[i].cells[1].text = b

spacer()
heading2("Step 4: Wire the DS18B20 Sensor (Temperature)")

table2 = doc.add_table(rows=4, cols=2)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "DS18B20 Pin"
hdr2[1].text = "ESP32 Pin"
rows2 = [("VCC", "3.3V"), ("GND", "GND"), ("DATA", "GPIO 4")]
for i, (a, b) in enumerate(rows2, 1):
    table2.rows[i].cells[0].text = a
    table2.rows[i].cells[1].text = b

spacer()
note("Place a 4.7kΩ resistor between the DS18B20 DATA pin and VCC pin. This is required — without it the sensor will not work.")

doc.add_page_break()

# ── PART 3 ────────────────────────────────────────────────────────────────────
heading1("PART 3 — Backend Setup")

heading2("Step 5: Install Python Dependencies")
body("Make sure Python 3.10 or higher is installed:")
code("python3 --version")
spacer()
body("Install the required packages:")
code("cd backend")
code("pip install -r requirements.txt")

spacer()
heading2("Step 6: Set Up PostgreSQL Database")
body("Install PostgreSQL if not installed:")
code("sudo apt install postgresql -y")
code("sudo systemctl start postgresql")

spacer()
body("Create the database and user:")
code('sudo -u postgres psql -c "CREATE DATABASE health_monitor;"')
code('sudo -u postgres psql -c "CREATE USER health_user WITH PASSWORD \'health_pass123\';"')
code('sudo -u postgres psql -c "GRANT ALL PRIVILEGES ON DATABASE health_monitor TO health_user;"')
code('sudo -u postgres psql -c "ALTER DATABASE health_monitor OWNER TO health_user;"')

spacer()
heading2("Step 7: Create the Database Tables")
code("cd backend")
code('python3 -c "from app.db.database import Base, engine; from app.models.vitals import Vital; Base.metadata.create_all(engine); print(\'Tables created\')"')
note("Run this as one single line in your terminal.")

spacer()
heading2("Step 8: Start the Backend Server")
code("cd backend")
code("uvicorn app.main:app --reload --host 0.0.0.0 --port 8000")
spacer()
body("Open your browser and go to:")
code("http://localhost:8000/docs")
body("You should see the Swagger API documentation page with 3 endpoints. This confirms the backend is running correctly.")

doc.add_page_break()

# ── PART 4 ────────────────────────────────────────────────────────────────────
heading1("PART 4 — Frontend Dashboard Setup")

heading2("Step 9: Install Node.js (if not installed)")
body("Check if Node.js is installed:")
code("node --version")
body("If not installed:")
code("sudo apt install nodejs npm -y    # Linux")
code("# Or download from https://nodejs.org  (Windows/macOS)")

spacer()
heading2("Step 10: Install Frontend Dependencies")
code("cd frontend")
code("npm install")

spacer()
heading2("Step 11: Start the Dashboard")
code("cd frontend")
code("npm run dev")
spacer()
body("Open your browser and go to:")
code("http://localhost:5173")
body("You will see the dashboard with Heart Rate, SpO2, and Temperature cards. They will show -- until the ESP32 starts sending data.")

doc.add_page_break()

# ── PART 5 ────────────────────────────────────────────────────────────────────
heading1("PART 5 — ESP32 Firmware Setup")

heading2("Step 12: Install PlatformIO")
body("Install PlatformIO CLI:")
code("pip install platformio")
spacer()
body("Verify installation:")
code("pio --version")

spacer()
heading2("Step 13: Find Your PC's IP Address")
body("You need your PC's local IP so the ESP32 can send data to it.")
code("hostname -I          # Linux/macOS")
code("ipconfig             # Windows")
body("Look for something like 192.168.1.x — that is your IP.")

spacer()
heading2("Step 14: Edit the Firmware Configuration")
body("Open the file:")
code("firmware/src/main.cpp")
spacer()
body("Find and edit these 3 lines at the top:")
code('const char* WIFI_SSID    = "YOUR_WIFI_NAME";')
code('const char* WIFI_PASS    = "YOUR_WIFI_PASSWORD";')
code('const char* API_ENDPOINT = "http://192.168.1.x:8000/api/vitals/";')
spacer()
note("Replace YOUR_WIFI_NAME and YOUR_WIFI_PASSWORD with your actual WiFi credentials. Replace 192.168.1.x with your actual PC IP from Step 13.")

spacer()
heading2("Step 15: Flash the ESP32")
body("Connect your ESP32 to your PC via USB cable, then run:")
code("cd firmware")
code("pio run --target upload")
note("First run will download the ESP32 toolchain and libraries automatically. This may take 5-10 minutes.")

spacer()
heading2("Step 16: Monitor Serial Output")
code("pio device monitor")
spacer()
body("You should see output like this:")
code("WiFi connected: 192.168.1.x")
code("HR: 72 bpm | SpO2: 98% | Temp: 36.50 C")
code("POST http://192.168.1.x:8000/api/vitals/ -> HTTP 201")
body("HTTP 201 means data was saved successfully.")

doc.add_page_break()

# ── PART 6 ────────────────────────────────────────────────────────────────────
heading1("PART 6 — Test Everything Together")

heading2("Step 17: Run All Three Components")
body("You need 2 terminals open at the same time:")

table3 = doc.add_table(rows=3, cols=2)
table3.style = "Table Grid"
hdr3 = table3.rows[0].cells
hdr3[0].text = "Terminal"
hdr3[1].text = "Command"
table3.rows[1].cells[0].text = "Terminal 1 (Backend)"
table3.rows[1].cells[1].text = "cd backend && uvicorn app.main:app --reload --host 0.0.0.0 --port 8000"
table3.rows[2].cells[0].text = "Terminal 2 (Frontend)"
table3.rows[2].cells[1].text = "cd frontend && npm run dev"

spacer()
body("Then make sure your ESP32 is powered on and connected to WiFi.")

spacer()
heading2("Step 18: Test Without ESP32 (Optional)")
body("Send a fake reading to verify the full stack works:")
code("curl -X POST http://localhost:8000/api/vitals/ \\")
code("  -H 'Content-Type: application/json' \\")
code("  -d '{\"patient_id\":1,\"heart_rate\":75,\"spo2\":98,\"temperature\":36.5}'")
spacer()
body("Then open http://localhost:5173 — you should see the values appear on the dashboard within 5 seconds.")

doc.add_page_break()

# ── Troubleshooting ───────────────────────────────────────────────────────────
heading1("Troubleshooting")

heading2("Backend won't start")
bullet("Make sure PostgreSQL is running: sudo systemctl start postgresql")
bullet("Make sure you ran the table creation command in Step 7")
bullet("Check the .env file in backend/ has the correct DATABASE_URL")

heading2("ESP32 won't connect to WiFi")
bullet("Double-check WIFI_SSID and WIFI_PASS in main.cpp")
bullet("Make sure your PC and ESP32 are on the same WiFi network")
bullet("Try moving the ESP32 closer to the router")

heading2("ESP32 connects but dashboard shows no data")
bullet("Check API_ENDPOINT IP matches your PC's IP (hostname -I)")
bullet("Make sure backend is running on port 8000")
bullet("Check serial monitor — it should show HTTP 201, not an error")

heading2("Dashboard not loading")
bullet("Make sure npm run dev is running in the frontend folder")
bullet("Open http://localhost:5173 (not 8000)")

spacer()

# ── Summary ───────────────────────────────────────────────────────────────────
heading1("Quick Reference")

table4 = doc.add_table(rows=6, cols=2)
table4.style = "Table Grid"
hdr4 = table4.rows[0].cells
hdr4[0].text = "Service"
hdr4[1].text = "URL / Command"
refs = [
    ("Backend API", "http://localhost:8000"),
    ("API Docs (Swagger)", "http://localhost:8000/docs"),
    ("Frontend Dashboard", "http://localhost:5173"),
    ("Database", "postgresql://health_user:health_pass123@localhost/health_monitor"),
    ("Serial Monitor", "pio device monitor (in firmware folder)"),
]
for i, (a, b) in enumerate(refs, 1):
    table4.rows[i].cells[0].text = a
    table4.rows[i].cells[1].text = b

spacer()
p = doc.add_paragraph("Good luck! — The system is designed to work out of the box once configured.")
p.runs[0].font.size = Pt(11)
p.runs[0].italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
path = "/home/anonymous_22/projects/health-monitor/Friend_Setup_Guide.docx"
doc.save(path)
print(f"Saved: {path}")
